from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
CAPTURE_DIR = ROOT / "docs" / "real-ui-captures"
MANIFEST_PATH = CAPTURE_DIR / "manifest.json"
OUTPUT_PATH = ROOT / "docs" / "psycare-real-ui-screenshots.docx"


ACCENT = "B00012"
DARK = "111827"
MUTED = "64748B"
LIGHT = "F8FAFC"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = BORDER) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:%s" % edge
        border = borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, bold=False, color=DARK, size=8.5) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(size)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor.from_string(DARK if level > 1 else ACCENT)


def add_small_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"].font.color.rgb = RGBColor.from_string(DARK)


def add_cover(document: Document, captures: list[dict]) -> None:
    add_heading(document, "PsyCare 2.0 Real UI Screenshot Capture", 0)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run(
        "Actual UI screenshots captured from the local running application."
    )
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    details = [
        ("Generated on", date.today().strftime("%d %B %Y")),
        ("Screens captured", str(len(captures))),
        ("Base URL", "http://localhost:8001"),
        (
            "Coverage goal",
            "Cover the main UI states used by the sequence diagrams across client, admin, and counsellor workflows.",
        ),
    ]
    for label, value in details:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = RGBColor.from_string(DARK)
        value_run = paragraph.add_run(value)
        value_run.font.size = Pt(10)
        value_run.font.color.rgb = RGBColor.from_string(MUTED)

    document.add_paragraph()
    add_small_note(
        document,
        "Note: the capture script only seeded browser-session values for language and client terms acceptance. The screens themselves come from the existing UI and in-app mock data.",
    )
    add_small_note(
        document,
        "The first screenshot intentionally shows the terms modal before acceptance; subsequent client screenshots show the accepted session state so the underlying pages are visible.",
    )
    document.add_page_break()


def add_coverage_table(document: Document, captures: list[dict]) -> None:
    add_heading(document, "Coverage Overview", 1)
    add_small_note(
        document,
        "Each screenshot is mapped to the sequence-diagram or module flow it helps demonstrate.",
    )

    for index, item in enumerate(captures, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.15)
        paragraph.paragraph_format.first_line_indent = Inches(-0.15)
        paragraph.paragraph_format.space_after = Pt(2)
        title_run = paragraph.add_run(f"{index:02d}. {item['title']} ")
        title_run.bold = True
        title_run.font.size = Pt(8.4)
        title_run.font.color.rgb = RGBColor.from_string(DARK)
        route_run = paragraph.add_run(f"({item['route']}) - ")
        route_run.font.size = Pt(8.2)
        route_run.font.color.rgb = RGBColor.from_string(MUTED)
        coverage_run = paragraph.add_run(item["coverage"])
        coverage_run.font.size = Pt(8.2)
        coverage_run.font.color.rgb = RGBColor.from_string(DARK)

    document.add_page_break()


def add_screenshot_pages(document: Document, captures: list[dict]) -> None:
    for index, item in enumerate(captures, start=1):
        add_heading(document, f"{index:02d}. {item['title']}", 1)

        route = document.add_paragraph()
        route.paragraph_format.space_after = Pt(2)
        route_run = route.add_run(f"Route: {item['route']}")
        route_run.font.size = Pt(8.5)
        route_run.font.bold = True
        route_run.font.color.rgb = RGBColor.from_string(MUTED)

        coverage = document.add_paragraph()
        coverage.paragraph_format.space_after = Pt(5)
        coverage_run = coverage.add_run(f"Coverage: {item['coverage']}")
        coverage_run.font.size = Pt(8.5)
        coverage_run.font.color.rgb = RGBColor.from_string(DARK)

        image_path = CAPTURE_DIR / item["file"]
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(8.65))

        if index != len(captures):
            document.add_page_break()


def main() -> None:
    captures = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    captures = [item for item in captures if item.get("file")]

    document = Document()
    configure_document(document)
    add_cover(document, captures)
    add_coverage_table(document, captures)
    add_screenshot_pages(document, captures)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
