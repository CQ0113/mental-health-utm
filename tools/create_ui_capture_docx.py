from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "ui-captures"
DOCX_PATH = ROOT / "docs" / "psycare-ui-feature-captures.docx"

W, H = 1600, 1000

COLORS = {
    "bg": "#eef4f7",
    "panel": "#ffffff",
    "line": "#d6e0e6",
    "text": "#1e293b",
    "muted": "#64748b",
    "navy": "#11394a",
    "teal": "#0f766e",
    "teal2": "#ccfbf1",
    "blue": "#2563eb",
    "blue2": "#dbeafe",
    "green": "#16a34a",
    "green2": "#dcfce7",
    "amber": "#d97706",
    "amber2": "#fef3c7",
    "red": "#dc2626",
    "red2": "#fee2e2",
    "purple": "#7c3aed",
    "purple2": "#ede9fe",
}


def font(size, bold=False):
    candidates = []
    if bold:
        candidates.extend([
            "C:/Windows/Fonts/arialbd.ttf",
            "C:/Windows/Fonts/segoeuib.ttf",
        ])
    candidates.extend([
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
    ])
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


F = {
    "title": font(30, True),
    "h1": font(24, True),
    "h2": font(20, True),
    "body": font(17),
    "small": font(14),
    "tiny": font(12),
    "bold": font(17, True),
    "button": font(15, True),
}


def hex_color(name):
    return COLORS.get(name, name)


def rounded(draw, xy, fill="panel", outline="line", radius=16, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=hex_color(fill), outline=hex_color(outline), width=width)


def text(draw, xy, value, fill="text", f=None, anchor=None):
    draw.text(xy, value, fill=hex_color(fill), font=f or F["body"], anchor=anchor)


def wrapped_text(draw, xy, value, max_chars, fill="text", f=None, line_gap=6):
    x, y = xy
    for line in wrap(value, width=max_chars):
        draw.text((x, y), line, fill=hex_color(fill), font=f or F["body"])
        y += (f or F["body"]).size + line_gap
    return y


def button(draw, xy, label, fill="teal", text_fill="#ffffff", w=150, h=42):
    x, y = xy
    rounded(draw, (x, y, x + w, y + h), fill=fill, outline=fill, radius=11)
    text(draw, (x + w / 2, y + h / 2), label, fill=text_fill, f=F["button"], anchor="mm")


def pill(draw, xy, label, fill="blue2", text_fill="blue", w=None):
    x, y = xy
    if w is None:
        w = max(84, len(label) * 8 + 24)
    rounded(draw, (x, y, x + w, y + 30), fill=fill, outline=fill, radius=15)
    text(draw, (x + w / 2, y + 15), label, fill=text_fill, f=F["small"], anchor="mm")


def field(draw, xy, label, value="", w=300, h=58, disabled=False):
    x, y = xy
    text(draw, (x, y), label, fill="muted", f=F["tiny"])
    fill = "#f8fafc" if disabled else "panel"
    rounded(draw, (x, y + 18, x + w, y + h), fill=fill, outline="line", radius=8)
    if value:
        text(draw, (x + 12, y + 32), value, fill="text", f=F["small"])


def table(draw, xy, columns, rows, col_widths, row_h=44, header_fill="#f1f5f9"):
    x, y = xy
    total_w = sum(col_widths)
    rounded(draw, (x, y, x + total_w, y + row_h * (len(rows) + 1)), fill="panel", outline="line", radius=10)
    draw.rectangle((x, y, x + total_w, y + row_h), fill=header_fill)
    cx = x
    for i, col in enumerate(columns):
        text(draw, (cx + 12, y + 13), col, fill="text", f=F["small"])
        cx += col_widths[i]
        draw.line((cx, y, cx, y + row_h * (len(rows) + 1)), fill=hex_color("line"))
    for idx, row in enumerate(rows):
        yy = y + row_h * (idx + 1)
        draw.line((x, yy, x + total_w, yy), fill=hex_color("line"))
        cx = x
        for i, cell in enumerate(row):
            text(draw, (cx + 12, yy + 13), str(cell), fill="text", f=F["small"])
            cx += col_widths[i]


def card(draw, xy, title, body=None, w=330, h=150, accent="teal"):
    x, y = xy
    rounded(draw, (x, y, x + w, y + h), fill="panel", outline="line", radius=14)
    draw.rectangle((x, y, x + 6, y + h), fill=hex_color(accent))
    text(draw, (x + 22, y + 18), title, fill="text", f=F["bold"])
    if body:
        wrapped_text(draw, (x + 22, y + 48), body, 34, fill="muted", f=F["small"], line_gap=4)


def shell(title, role="Admin"):
    img = Image.new("RGB", (W, H), hex_color("bg"))
    d = ImageDraw.Draw(img)
    rounded(d, (34, 28, 246, 970), fill="navy", outline="navy", radius=18)
    text(d, (64, 66), "PsyCare 2.0", fill="#ffffff", f=F["h2"])
    menu = ["Dashboard", "Appointments", "Clients", "Resources", "Reports", "Settings"]
    y = 136
    for item in menu:
        fill = "#d1fae5" if item in title or (title.startswith("Dashboard") and item == "Dashboard") else "#d9eef2"
        text(d, (64, y), item, fill=fill, f=F["small"])
        y += 48
    rounded(d, (276, 28, 1566, 92), fill="panel", outline="line", radius=18)
    text(d, (304, 50), title, fill="text", f=F["h1"])
    pill(d, (1330, 48), role, fill="teal2", text_fill="teal", w=110)
    text(d, (1452, 50), "Aina Admin", fill="muted", f=F["small"])
    return img, d


def save(img, name):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / f"{name}.png"
    img.save(path, "PNG")
    return path


def screen_um_counsellor():
    img, d = shell("Counsellor PPsi Management", "Admin")
    card(d, (286, 118), "UM01 Onboard Counsellor", "List, search, add, validate duplicate PPsi, and confirm save.", 430, 122)
    field(d, (744, 124), "Search counsellor", "Dr. Nadia / PPsi-1187", 330)
    button(d, (1106, 142), "+ Add Counsellor", "teal", w=180)
    rows = [
        ["Dr. Nadia Rahman", "PPsi-1187", "UTM KL", "Active"],
        ["Mr. Daniel Lee", "PPsi-0921", "JB Campus", "Active"],
        ["Ms. Sara Tan", "PPsi-1045", "JB Campus", "Inactive"],
    ]
    table(d, (286, 276), ["Name", "PPsi No", "Location", "Status"], rows, [360, 220, 220, 150])
    card(d, (286, 520), "Add Counsellor Form", "Name, email, PPsi number, worker number, location, phone, status, service dates.", 520, 260, "blue")
    field(d, (836, 520), "PPsi No", "PPsi-1187", 240)
    field(d, (1102, 520), "Worker No", "WK-2238", 240)
    field(d, (836, 604), "Email", "nadia.rahman@utm.my", 342)
    field(d, (1204, 604), "Location", "UTM JB Counselling Room", 300)
    button(d, (836, 708), "Confirm Save", "teal", w=150)
    rounded(d, (286, 824, 1504, 906), fill="red2", outline="red", radius=12)
    text(d, (314, 850), "Duplicate record warning", fill="red", f=F["bold"])
    text(d, (314, 878), "PPsi number already exists. Admin must correct the value before saving.", fill="text", f=F["small"])
    return save(img, "01-user-management-counsellor")


def screen_um_client_profile():
    img, d = shell("Client Profile Search And Details", "Admin")
    field(d, (286, 124), "Search by name / matrix / email", "Aina Tan / A23CS0112", 430)
    button(d, (742, 142), "Search", "teal", w=110)
    rows = [
        ["Aina Tan", "A23CS0112", "Student", "Profile locked"],
        ["Amir Lim", "W9872", "Staff", "Review required"],
        ["Nur Izzah", "A22CS0999", "Student", "Active"],
    ]
    table(d, (286, 224), ["Client", "ID No", "Type", "Status"], rows, [280, 220, 180, 250])
    card(d, (286, 456), "UM02 Client Details", "Read-only sections for profile, appointment history, and declaration history.", 420, 128)
    rows2 = [
        ["Profile", "Faculty: Computing, Phone: 012-222 7844"],
        ["Appointment", "AS-2026-0042, Pending counsellor review"],
        ["Declaration", "Submitted, awaiting verification"],
    ]
    table(d, (740, 456), ["Section", "Sample Information"], rows2, [190, 540])
    rounded(d, (286, 742, 1504, 828), fill="amber2", outline="amber", radius=12)
    text(d, (314, 768), "Access denied state", fill="amber", f=F["bold"])
    text(d, (314, 796), "Shown when a counsellor attempts to open a client outside their permitted caseload.", fill="text", f=F["small"])
    return save(img, "02-user-management-client-profile")


def screen_as_booking():
    img, d = shell("Smart Appointment Form", "Client")
    card(d, (286, 118), "AS01 / AS03 Book Appointment", "Applicant info, calendar, available slot, appointment details, declaration, and summary.", 470, 128)
    field(d, (786, 124), "Reference No", "AS-2026-0084", 220, disabled=True)
    field(d, (1036, 124), "Client", "Aina Tan", 240, disabled=True)
    field(d, (1286, 124), "Session Type", "Online", 220)
    rounded(d, (286, 284, 640, 770), fill="panel", outline="line", radius=16)
    text(d, (316, 316), "June 2026", f=F["h2"])
    days = ["Mon", "Tue", "Wed", "Thu", "Fri"]
    x = 316
    for day in days:
        text(d, (x, 368), day, fill="muted", f=F["small"])
        x += 62
    for i in range(1, 21):
        xx = 316 + ((i - 1) % 5) * 62
        yy = 410 + ((i - 1) // 5) * 58
        fill = "teal2" if i in [8, 15] else "#f8fafc"
        rounded(d, (xx, yy, xx + 44, yy + 38), fill=fill, outline="line", radius=8)
        text(d, (xx + 22, yy + 19), str(i), fill="teal" if i in [8, 15] else "text", f=F["small"], anchor="mm")
    rows = [
        ["10:00", "Dr. Nadia", "Online", "Available"],
        ["14:00", "Mr. Daniel", "Physical", "Available"],
        ["16:00", "Dr. Nadia", "Online", "Full"],
    ]
    table(d, (690, 284), ["Time", "Counsellor", "Type", "Status"], rows, [130, 220, 160, 150])
    field(d, (690, 538), "Issue Summary", "Stress, sleep difficulty, and assignment anxiety", 610, 66)
    field(d, (690, 626), "Attachment", "medical_note.pdf", 300)
    rounded(d, (690, 720, 1504, 828), fill="green2", outline="green", radius=12)
    text(d, (718, 748), "Booking summary", fill="green", f=F["bold"])
    text(d, (718, 778), "Appointment submitted as pending. Declaration record created. Meeting link generated for online session.", fill="text", f=F["small"])
    return save(img, "03-appointment-booking")


def screen_as_followup():
    img, d = shell("Follow-Up Appointment", "Client")
    rows = [
        ["AS-2026-0031", "Completed", "Eligible", "Dr. Nadia"],
        ["AS-2026-0020", "Cancelled", "Not eligible", "Mr. Daniel"],
        ["AS-2026-0017", "Completed", "Eligible", "Dr. Nadia"],
    ]
    table(d, (286, 142), ["Reference", "Status", "Follow-up", "Counsellor"], rows, [220, 180, 210, 240])
    card(d, (286, 392), "AS02 Follow-Up Flow", "Select eligible appointment, lock previous reference, then choose new date and slot.", 420, 142)
    field(d, (740, 402), "Previous Appointment", "AS-2026-0031", 260, disabled=True)
    field(d, (1024, 402), "New Slot", "18 Jun 2026, 10:00", 300)
    field(d, (740, 504), "Session Type", "Online", 220)
    button(d, (1024, 522), "Submit Follow-Up", "teal", w=180)
    rounded(d, (286, 706, 760, 824), fill="amber2", outline="amber", radius=12)
    text(d, (314, 734), "Not eligible state", fill="amber", f=F["bold"])
    text(d, (314, 764), "Show when selected appointment is not eligible for follow-up.", fill="text", f=F["small"])
    rounded(d, (806, 706, 1504, 824), fill="red2", outline="red", radius=12)
    text(d, (834, 734), "No available follow-up slot", fill="red", f=F["bold"])
    text(d, (834, 764), "Show when selected date or counsellor has no available follow-up slot.", fill="text", f=F["small"])
    return save(img, "04-appointment-followup")


def screen_as_slot_manager():
    img, d = shell("Slot Manager", "Admin")
    card(d, (286, 118), "AS04 / AS05 / AS06 Manage Slots", "Manual slot, bulk generation, CSV import, replace existing dates, and save draft changes.", 540, 128)
    field(d, (858, 124), "Counsellor", "Dr. Nadia Rahman", 260)
    field(d, (1136, 124), "Date Range", "10 Jun - 30 Jun", 260)
    button(d, (1420, 142), "Save Changes", "teal", w=150)
    rows = [
        ["12 Jun", "10:00-11:00", "Online", "Draft"],
        ["12 Jun", "14:00-15:00", "Physical", "Saved"],
        ["18 Jun", "09:00-10:00", "Online", "Remove"],
    ]
    table(d, (286, 286), ["Date", "Time", "Session", "Draft Status"], rows, [160, 220, 170, 180])
    card(d, (286, 548), "Manual Add", "Date, start/end time, counsellor, location, session type. Shows overlap error when invalid.", 350, 180, "blue")
    card(d, (666, 548), "Bulk Generate", "Weekdays, date range, session types, replace existing option, generation summary.", 350, 180, "purple")
    card(d, (1046, 548), "CSV Import", "CSV upload, template validation, valid rows, skipped rows, import summary.", 350, 180, "amber")
    rounded(d, (286, 796, 1504, 884), fill="green2", outline="green", radius=12)
    text(d, (314, 824), "Mocked import summary", fill="green", f=F["bold"])
    text(d, (314, 854), "42 rows processed, 38 valid, 4 skipped. 6 saved slots marked for removal after confirmation.", fill="text", f=F["small"])
    return save(img, "05-slot-manager")


def screen_as_queue():
    img, d = shell("Appointment Queue And Verification", "Admin")
    rows = [
        ["AS-2026-0084", "Aina Tan", "Needs review", "Online"],
        ["AS-2026-0085", "Amir Lim", "Counsellor review", "Physical"],
        ["AS-2026-0086", "Nur Izzah", "Pending", "Online"],
    ]
    table(d, (286, 132), ["Reference", "Client", "Status", "Session"], rows, [230, 250, 250, 160])
    rounded(d, (286, 394, 1504, 802), fill="panel", outline="line", radius=16)
    text(d, (318, 424), "Selected Appointment Details", f=F["h2"])
    field(d, (318, 474), "Client", "Aina Tan", 250, disabled=True)
    field(d, (594, 474), "Slot", "15 Jun 2026, 10:00", 280, disabled=True)
    field(d, (904, 474), "Counsellor", "Dr. Nadia Rahman", 300, disabled=True)
    field(d, (318, 584), "Admin Review Note", "Information is complete. Forward to counsellor.", 540)
    field(d, (884, 584), "Counsellor Review Note", "Approved for first session.", 470)
    button(d, (318, 704), "Approve For Review", "blue", w=190)
    button(d, (538, 704), "Approve Appointment", "teal", w=190)
    rounded(d, (770, 704, 1450, 764), fill="red2", outline="red", radius=12)
    text(d, (792, 724), "Missing review info error appears before status update.", fill="red", f=F["small"])
    return save(img, "06-appointment-verification")


def screen_ta_attendance():
    img, d = shell("Telemedicine And Attendance", "Counsellor")
    card(d, (286, 118), "TA01 Join Online Session", "Appointment detail page validates access and meeting link before opening online session.", 480, 132)
    field(d, (806, 126), "Appointment Ref", "AS-2026-0084", 250)
    button(d, (1084, 144), "Join Online Session", "teal", w=190)
    rows = [
        ["Aina Tan", "Present", "online_auto", "10:03"],
        ["Amir Lim", "Absent", "manual", "-"],
        ["Nur Izzah", "Present", "physical_qr", "14:10"],
    ]
    table(d, (286, 314), ["Participant", "Status", "Method", "Checked In"], rows, [280, 170, 200, 170])
    card(d, (286, 572), "TA02 Manual Attendance", "Counsellor selects attendance status and saves participant record.", 380, 160, "blue")
    card(d, (700, 572), "TA03 Physical QR", "Client scans QR, confirms identity, and sees attendance confirmation.", 380, 160, "green")
    card(d, (1114, 572), "TA04 Online Auto-Log", "System detects verified online join event and auto-records attendance.", 380, 160, "purple")
    rounded(d, (286, 820, 1504, 902), fill="red2", outline="red", radius=12)
    text(d, (314, 850), "Required exception states", fill="red", f=F["bold"])
    text(d, (314, 878), "Invalid QR, duplicate check-in, unauthorized scan, non-physical session, invalid meeting link, and unavailable session.", f=F["small"])
    return save(img, "07-attendance")


def screen_declaration():
    img, d = shell("Declaration Form And Review", "Client")
    rounded(d, (286, 126, 760, 790), fill="panel", outline="line", radius=16)
    text(d, (318, 158), "Client Declaration", f=F["h2"])
    field(d, (318, 214), "Client", "Aina Tan", 360, disabled=True)
    field(d, (318, 300), "Appointment", "AS-2026-0084", 360, disabled=True)
    rounded(d, (318, 398, 704, 506), fill="#f8fafc", outline="line", radius=10)
    wrapped_text(d, (342, 424), "I confirm that the submitted profile and appointment information is true and complete.", 44, f=F["small"])
    text(d, (342, 548), "[x] I agree to submit this declaration.", f=F["bold"])
    button(d, (342, 604), "Submit Declaration", "teal", w=190)
    rounded(d, (824, 126, 1504, 790), fill="panel", outline="line", radius=16)
    text(d, (856, 158), "Admin / Counsellor Review", f=F["h2"])
    rows = [
        ["Status", "Submitted"],
        ["Submitted At", "06 Jun 2026, 09:20"],
        ["Verifier", "Dr. Nadia Rahman"],
    ]
    table(d, (856, 224), ["Item", "Value"], rows, [190, 330])
    field(d, (856, 474), "Correction Note", "Please update phone number before verification.", 500)
    button(d, (856, 594), "Verify", "green", w=120)
    button(d, (1000, 594), "Require Correction", "amber", w=190)
    rounded(d, (286, 842, 1504, 918), fill="amber2", outline="amber", radius=12)
    text(d, (314, 866), "Capture states: checkbox required, declaration unavailable, already verified, incomplete declaration, verification success.", f=F["small"])
    return save(img, "08-declaration")


def screen_ct_dashboard():
    img, d = shell("Dashboard Emotion Tracker", "Client")
    field(d, (286, 126), "Emotion Date", "06 Jun 2026", 220)
    field(d, (534, 126), "Emotion Score", "7 / 10", 180)
    button(d, (734, 144), "Save Emotion", "teal", w=150)
    rounded(d, (286, 260, 894, 700), fill="panel", outline="line", radius=16)
    text(d, (316, 292), "Emotion Scores Graph", f=F["h2"])
    points = [(360, 610), (470, 520), (580, 550), (690, 420), (800, 380)]
    for i in range(len(points) - 1):
        d.line((points[i], points[i + 1]), fill=hex_color("teal"), width=5)
    for p in points:
        d.ellipse((p[0] - 8, p[1] - 8, p[0] + 8, p[1] + 8), fill=hex_color("teal"))
    text(d, (350, 642), "Mon", f=F["tiny"], fill="muted")
    text(d, (460, 642), "Tue", f=F["tiny"], fill="muted")
    text(d, (570, 642), "Wed", f=F["tiny"], fill="muted")
    text(d, (680, 642), "Thu", f=F["tiny"], fill="muted")
    text(d, (790, 642), "Fri", f=F["tiny"], fill="muted")
    card(d, (940, 260), "CT01 Save State", "Success message: You logged today's emotion score. Meaningful words: steady, brave, aware.", 460, 148, "green")
    card(d, (940, 444), "CT03 Filter State", "Date range filter shows client or counsellor emotion history graph.", 460, 148, "blue")
    rounded(d, (940, 628, 1400, 730), fill="red2", outline="red", radius=12)
    text(d, (966, 656), "Errors", fill="red", f=F["bold"])
    text(d, (966, 684), "Invalid score, future date, or no emotion history found.", f=F["small"])
    return save(img, "09-emotion-dashboard")


def screen_ct_chatbot():
    img, d = shell("AI Counsellor Chatbot", "Client")
    rounded(d, (286, 118, 1050, 882), fill="panel", outline="line", radius=18)
    text(d, (320, 152), "AI Counsellor Chat", f=F["h2"])
    bubbles = [
        ("AI", "Hi Aina, how are you feeling today?", "#f1f5f9"),
        ("Client", "I feel overwhelmed by assignments.", "#ccfbf1"),
        ("AI", "That sounds heavy. Try one small step and a breathing pause.", "#f1f5f9"),
    ]
    y = 220
    for sender, msg, fill in bubbles:
        x = 330 if sender == "AI" else 560
        rounded(d, (x, y, x + 430, y + 78), fill=fill, outline="line", radius=18)
        text(d, (x + 18, y + 14), sender, fill="muted", f=F["tiny"])
        wrapped_text(d, (x + 18, y + 36), msg, 48, f=F["small"])
        y += 110
    field(d, (330, 680), "Message", "I need help calming down.", 530)
    button(d, (880, 698), "Send", "teal", w=100)
    button(d, (330, 788), "Save Chat", "blue", w=120)
    card(d, (1110, 142), "CT02 Risk Detected", "High-stress guidance message appears and conversation can be marked for counsellor review.", 390, 170, "red")
    card(d, (1110, 358), "Quick Replies", "I feel stressed, I need grounding, I want to talk to someone.", 390, 140, "purple")
    card(d, (1110, 540), "Saved Chat", "The chat session is saved for counsellor review.", 390, 140, "green")
    return save(img, "10-ai-chatbot")


def screen_ct_caseload():
    img, d = shell("Caseload And Flagged Client", "Counsellor")
    rows = [
        ["Aina Tan", "High", "AI chatbot", "Open"],
        ["Amir Lim", "Moderate", "Emotion", "In review"],
        ["Nur Izzah", "High", "Psychometric", "Open"],
    ]
    table(d, (286, 132), ["Client", "Risk", "Source", "Status"], rows, [240, 160, 220, 170])
    rounded(d, (286, 394, 1504, 834), fill="panel", outline="line", radius=16)
    text(d, (318, 426), "Flagged Client Detail Panel", f=F["h2"])
    field(d, (318, 484), "Selected Client", "Aina Tan", 250, disabled=True)
    field(d, (596, 484), "Latest Risk", "High stress from chatbot", 350, disabled=True)
    field(d, (974, 484), "Last Appointment", "AS-2026-0031 Completed", 350, disabled=True)
    field(d, (318, 592), "Review Note", "Client needs follow-up check within 3 days.", 520)
    field(d, (870, 592), "Task Title", "Schedule intervention call", 380)
    button(d, (318, 714), "Save Review", "blue", w=150)
    button(d, (496, 714), "Create Task", "teal", w=140)
    rounded(d, (688, 710, 1406, 770), fill="green2", outline="green", radius=12)
    text(d, (710, 730), "Show review saved success and task created success after valid details.", f=F["small"])
    return save(img, "11-caseload")


def screen_er_resource():
    img, d = shell("Resource Library", "Client")
    card(d, (286, 118), "ER02 Access Learning Materials", "Search, filter, open resource, and record access log.", 430, 128)
    field(d, (746, 124), "Keyword", "stress", 250)
    field(d, (1016, 124), "Category", "Stress Management", 250)
    button(d, (1290, 142), "Filter", "teal", w=110)
    card(d, (286, 286), "Breathing Toolkit", "Category: Stress. Type: Toolkit. Duration: 8 minutes. Open resource URL.", 340, 180, "teal")
    card(d, (666, 286), "Sleep Reset Guide", "Category: Sleep. Type: Article. Duration: 5 minutes.", 340, 180, "blue")
    card(d, (1046, 286), "Anxiety Video", "Category: Anxiety. Type: Video. Duration: 12 minutes.", 340, 180, "purple")
    rounded(d, (286, 566, 1504, 850), fill="panel", outline="line", radius=16)
    text(d, (318, 598), "ER01 Admin Upload Learning Material", f=F["h2"])
    field(d, (318, 654), "Title", "Managing Exam Stress", 330)
    field(d, (672, 654), "URL", "https://resources.utm.my/stress-guide", 430)
    field(d, (1130, 654), "Visibility", "Published", 220)
    button(d, (318, 762), "Upload", "teal", w=120)
    text(d, (480, 774), "States: missing title/URL, invalid resource URL, upload success, no matching resource.", f=F["small"], fill="muted")
    return save(img, "12-resource-library")


def screen_ps_forum():
    img, d = shell("Peer Support Forum", "Client")
    rounded(d, (286, 118, 780, 790), fill="panel", outline="line", radius=16)
    text(d, (318, 152), "PS01 Submit Forum Post", f=F["h2"])
    field(d, (318, 218), "Title", "How do you manage exam stress?", 380)
    field(d, (318, 304), "Category", "Study Stress", 260)
    rounded(d, (318, 408, 720, 566), fill="#f8fafc", outline="line", radius=10)
    wrapped_text(d, (342, 430), "I am feeling nervous before finals. What helps you stay focused?", 50, f=F["small"])
    button(d, (318, 620), "Submit Post", "teal", w=140)
    rounded(d, (318, 696, 720, 754), fill="amber2", outline="amber", radius=10)
    text(d, (340, 716), "Unsafe content is queued for review.", f=F["small"], fill="amber")
    rounded(d, (828, 118, 1504, 790), fill="panel", outline="line", radius=16)
    text(d, (860, 152), "PS02 Moderate Forum", f=F["h2"])
    rows = [
        ["Post 104", "Pending review", "Low safety", "Hide"],
        ["Post 105", "Published", "Safe", "Approve"],
        ["Post 106", "Hidden", "Reported", "Restore"],
    ]
    table(d, (860, 220), ["Post", "Status", "Signal", "Action"], rows, [150, 180, 170, 130])
    field(d, (860, 488), "Moderation Reason", "Contains unsafe advice.", 440)
    button(d, (860, 602), "Confirm Action", "red", w=170)
    return save(img, "13-forum")


def screen_sa_psychometric():
    img, d = shell("Psychometric Self-Assessment", "Client")
    rounded(d, (286, 118, 760, 820), fill="panel", outline="line", radius=16)
    text(d, (318, 152), "SA01 Take Test", f=F["h2"])
    card(d, (318, 210), "DASS-21", "Published test. Estimated 10 minutes.", 380, 120, "teal")
    text(d, (318, 372), "Question 3 of 21", f=F["bold"])
    text(d, (318, 420), "I found it hard to wind down.", f=F["small"])
    opts = ["0 Never", "1 Sometimes", "2 Often", "3 Almost always"]
    y = 470
    for opt in opts:
        text(d, (344, y), "( ) " + opt, f=F["small"])
        y += 42
    button(d, (318, 678), "Submit / Review", "teal", w=160)
    rounded(d, (824, 118, 1504, 820), fill="panel", outline="line", radius=16)
    text(d, (856, 152), "Result And Triage", f=F["h2"])
    card(d, (856, 214), "Result Summary", "Score: 72%. Risk level: High. Recommendation shown to client.", 460, 150, "red")
    rows = [
        ["Aina Tan", "DASS-21", "High", "Open"],
        ["Amir Lim", "PHQ-9", "Moderate", "Reviewed"],
        ["Nur Izzah", "GAD-7", "Low", "Closed"],
    ]
    table(d, (856, 438), ["Client", "Test", "Risk", "Status"], rows, [180, 130, 130, 140])
    card(d, (856, 690), "SA03 Manage Test", "Admin uploads PDF, validates file type, generates questions and options.", 460, 110, "blue")
    return save(img, "14-psychometric")


def screen_ui_summary():
    img, d = shell("UI Design Coverage Summary", "Admin")
    card(d, (286, 118), "Purpose", "Mock screenshots cover the main UI flows from sequence diagrams using realistic sample records.", 470, 120)
    rows = [
        ["User Management", "Counsellor onboarding, profile search, profile detail"],
        ["Appointment", "Booking, follow-up, slots, appointment verification"],
        ["Attendance", "Manual, QR, online auto-log"],
        ["Declaration", "Client submit and staff verification"],
        ["Tracking", "Emotion, chatbot, caseload, task"],
        ["Resources", "Admin upload and client access"],
        ["Forum", "Post submission and moderation"],
        ["Psychometric", "Test, results, triage, test generation"],
    ]
    table(d, (286, 286), ["Module", "Captured UI Flow"], rows, [280, 720], row_h=52)
    rounded(d, (286, 842, 1504, 918), fill="blue2", outline="blue", radius=12)
    text(d, (314, 866), "Mocked data used: sample clients, appointment references, counsellors, slots, risk flags, resources, forum posts, and test submissions.", f=F["small"])
    return save(img, "00-coverage-summary")


SCREENS = [
    ("Coverage Summary", "All modules", screen_ui_summary),
    ("User Management - Counsellor Onboarding", "UM01", screen_um_counsellor),
    ("User Management - Client Profile Search", "UM02 / UM03", screen_um_client_profile),
    ("Appointment Booking", "AS01 / AS03", screen_as_booking),
    ("Follow-Up Appointment", "AS02", screen_as_followup),
    ("Slot Manager", "AS04 / AS05 / AS06", screen_as_slot_manager),
    ("Appointment Verification Queue", "AS07", screen_as_queue),
    ("Telemedicine And Attendance", "TA01 / TA02 / TA03 / TA04", screen_ta_attendance),
    ("Declaration Form And Review", "DC01 / DC02 / DC03", screen_declaration),
    ("Emotion Tracker And History", "CT01 / CT03", screen_ct_dashboard),
    ("AI Counsellor Chatbot", "CT02", screen_ct_chatbot),
    ("Caseload And Flagged Client", "CT04", screen_ct_caseload),
    ("Resource Library", "ER01 / ER02", screen_er_resource),
    ("Peer Support Forum", "PS01 / PS02", screen_ps_forum),
    ("Psychometric Self-Assessment", "SA01 / SA02 / SA03", screen_sa_psychometric),
]


def set_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)


def build_docx(paths):
    doc = Document()
    set_landscape(doc.sections[0])

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(24)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(17, 57, 74)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PsyCare 2.0 UI Feature Capture Pack")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(17, 57, 74)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mocked UI screenshots prepared from the SRS sequence-flow coverage checklist.")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph("Included mocked data: clients, counsellors, appointment references, slots, attendance records, declarations, risk flags, resources, forum posts, and psychometric submissions.")
    doc.add_paragraph("Use this document as a visual checklist for UI design screens and key states.")

    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = "Table Grid"
    header = summary_table.rows[0].cells
    header[0].text = "No"
    header[1].text = "Screen"
    header[2].text = "Use Case Coverage"
    for i, (title, coverage, _) in enumerate(SCREENS, start=1):
        row = summary_table.add_row().cells
        row[0].text = str(i)
        row[1].text = title
        row[2].text = coverage

    doc.add_page_break()

    for idx, ((title, coverage, _), path) in enumerate(zip(SCREENS, paths), start=1):
        doc.add_picture(str(path), width=Inches(10.85))
        if idx != len(paths):
            doc.add_page_break()

    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    paths = []
    for _, _, fn in SCREENS:
        paths.append(fn())
    build_docx(paths)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
